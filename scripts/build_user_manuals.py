"""
Generate three audience-specific Word manuals for the Secondary Targets & PBIS
module:

  - User_Manual_Admin.docx
  - User_Manual_Sales_User.docx
  - User_Manual_Hierarchy_User.docx

Run from the repo root:
    python3 scripts/build_user_manuals.py
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

OUT_DIR = Path(__file__).resolve().parent.parent

PRIMARY_BLUE = RGBColor(0x16, 0x32, 0x5C)
ACCENT_GREEN = RGBColor(0x2E, 0x84, 0x4A)


def style(doc):
    """Apply consistent base styling."""
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def add_cover(doc, title, audience):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Elite Foods")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = ACCENT_GREEN

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = PRIMARY_BLUE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"User Manual — {audience}")
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Secondary Targets & PBIS (Performance-Based Incentive Scheme)")
    run.italic = True

    doc.add_paragraph()


def h(doc, text, level=1):
    return doc.add_heading(text, level=level)


def para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h_ in enumerate(headers):
        hdr[i].text = h_
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = "" if v is None else str(v)
    doc.add_paragraph()


def callout(doc, text):
    """A subtle 'note' paragraph."""
    p = doc.add_paragraph()
    run = p.add_run("Note — " + text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x70, 0x6E, 0x6B)


# ============================================================================
# ADMIN MANUAL
# ============================================================================

def build_admin():
    doc = Document()
    style(doc)
    add_cover(doc, "Secondary Targets & PBIS", "Administrator")

    h(doc, "1. Introduction", 1)
    para(doc,
        "This manual is for the administrator who configures and operates the "
        "Secondary Targets and PBIS module — defining what is measured, setting "
        "incentive slabs, loading targets, running the monthly compute, and "
        "troubleshooting issues raised by sales staff or managers."
    )

    h(doc, "2. What the Module Does", 1)
    para(doc,
        "The module digitises three things the MIS team used to maintain in Excel."
    )
    table(doc,
        ["BRD Reference", "Capability", "Owner"],
        [
            ["§3",   "Secondary Targets vs Achievement engine",  "Admin sets up; daily batch runs"],
            ["§3.4", "Focus Pack masters (channel-scoped SKUs)", "Admin"],
            ["§5",   "Monthly Performance-Based Incentive Scheme","Admin runs monthly"],
            ["§4",   "KPI visibility for sales staff & managers", "Everyone"],
        ],
    )

    h(doc, "3. Key Concepts & Glossary", 1)
    table(doc,
        ["Term", "Meaning"],
        [
            ["DSM / SSA", "Field staff whose secondary sales are tracked. Identified by the DSM or SSA Salesforce profile."],
            ["Target Criterion", "A reusable parameter definition (Revenue, FP ECO, TLSD, …). Stored on Target_Criteria__c. Configured once."],
            ["Focus Pack", "A bundle of SKUs with a minimum qualifying quantity. Used by FOCUS_PACK_* criteria. Stored on Focused_Pack__c."],
            ["Secondary Target", "One DSM/SSA's target for one criterion over one date range. Stored on Secondary_Target__c, name format STGT-0001."],
            ["Achievement", "Actual measured value, computed daily by the Secondary Achievement Batch."],
            ["Working Days", "Days the DSM/SSA had Start Day done; divisor for daily-averaging criteria (TC / PC / TLSD)."],
            ["Incentive Slab", "A payout bracket on Incentive_Slab__c. Per criterion, per channel, optionally per Focus Pack."],
            ["Incentive Credit", "One earned payout row on Incentive_Credit__c (Source = 'Secondary PBIS'). Monthly payout = SUM of these rows."],
            ["PBIS Run", "Monthly compute that converts achievements into Incentive_Credit__c rows."],
            ["Period", "Year + Month. All PBIS results are keyed by these two fields."],
        ],
    )

    h(doc, "4. Tabs You Will Use", 1)
    table(doc,
        ["Tab", "Purpose"],
        [
            ["Target Criteria Manager",  "Define / edit reusable criteria (4-step wizard)."],
            ["Focused Pack",             "Create / edit Focus Pack masters."],
            ["Secondary Target Manager", "Create / edit targets per user / criterion / period. Bulk CSV import / export. Per-target recalc and View Calculation."],
            ["Secondary PBIS Slabs",     "Maintain payout slabs per criterion / channel / Focus Pack. Bulk add supported."],
            ["Secondary PBIS Console",   "Run PBIS for a selected month, review per-user totals, drill into the breakdown, export CSV."],
            ["Secondary KPI Dashboard",  "Mobile-friendly KPI dashboard used by sales staff and managers (read-only for them)."],
        ],
    )

    h(doc, "5. Setup Order (One-Time)", 1)
    para(doc, "Run this sequence in order — later steps depend on master data from earlier steps.")
    numbered(doc, [
        "Create Target Criteria — every parameter the business wants to track (Revenue, TC, PC, UBO/ECO, TLSD, FP Revenue, FP ECO).",
        "Create Focus Packs — SKU bundles + minimum qualifying quantities used by FOCUS_PACK_* criteria.",
        "Create Secondary Targets — assign each DSM / SSA the criteria they will be measured on with target value and active date range. Use the CSV importer for bulk uploads.",
        "Configure Incentive Slabs — payout brackets per criterion / channel / Focus Pack.",
        "Schedule the daily achievement batch so achievements stay current.",
        "Schedule the monthly PBIS batch, or run on demand from the PBIS Console.",
    ])

    h(doc, "6. Defining a Target Criterion", 1)
    para(doc, "Open Target Criteria Manager and click New Criterion. The wizard has four steps:")
    table(doc,
        ["Step", "What you set"],
        [
            ["1. Basics",     "Name, Active flag."],
            ["2. Calculation","Operator (see Appendix A) and the operator-specific fields. Fields adapt to the operator so irrelevant inputs are hidden."],
            ["3. Source",     "Object, User Field, Date Field, optional Filters / Filter Logic, optional Secondary Source Objects (UI builder — no JSON typing)."],
            ["4. Preview",    "Sample run for a date range so you can sanity-check the configuration before saving."],
        ],
    )

    h(doc, "Operator quick-pick", 2)
    table(doc,
        ["Sales parameter (BRD)", "Operator", "Distinct / Numerator / Denominator"],
        [
            ["Cumulative Secondary Revenue", "SUM",                 "—"],
            ["UBO / ECO",                    "COUNT_DISTINCT",      "Distinct = outlet field"],
            ["Total Calls (TC)",             "DAILY_UNIQUE_AVG",    "Distinct = outlet; add Secondary Source Objects to union order + visit-form sources"],
            ["Productive Calls (PC)",        "DAILY_UNIQUE_AVG",    "Distinct = outlet, orders only"],
            ["Average TLSD",                 "DAILY_RATIO_AVG",     "Numerator = line items, Denominator = orders"],
            ["Focus Pack Revenue",           "FOCUS_PACK_REVENUE",  "Focus Pack is set on each target row"],
            ["Focus Pack ECO",               "FOCUS_PACK_ECO",      "Focus Pack on the target; SUM(pack qty) >= Min Qty within one order"],
        ],
    )
    callout(doc, "Tick Use Attendance Divisor on daily-averaging criteria so the "
                 "divisor is working days (BRD §3.3.1) instead of calendar days.")

    h(doc, "7. Creating Focus Packs", 1)
    para(doc, "Open Focused Pack and click New. Set:")
    bullets(doc, [
        "Name — descriptive (e.g. 'Diwali Focus Pack — CPD').",
        "Sales Channel — drives the SKU picker.",
        "Minimum Qualifying Quantity — total units across pack SKUs needed in ONE order to count as Focus Pack ECO.",
        "Items — pick SKUs from the channel-filtered list, each with a minimum quantity.",
    ])

    h(doc, "8. Creating Secondary Targets", 1)

    h(doc, "8.1 Manual creation", 2)
    numbered(doc, [
        "Open Secondary Target Manager and click New.",
        "Pick User (DSM/SSA), Target Criterion, Focus Pack (only for FOCUS_PACK_* criteria), Sales Channel, Year, Start Date, End Date and Target Value.",
        "Save.",
    ])
    para(doc, "The system enforces:")
    bullets(doc, [
        "End Date >= Start Date always.",
        "End Date >= today on create (Start Date can be in the past so quarter / year targets work).",
        "No duplicate active overlapping target for the same (User, Criterion, Focus Pack) combination.",
    ])

    h(doc, "8.2 Bulk import via CSV", 2)
    numbered(doc, [
        "Click Download CSV Template for a header row with a sample line.",
        "Fill the spreadsheet with the required columns.",
        "Click Import CSV and pick the file.",
    ])
    para(doc, "Required columns:")
    bullets(doc, [
        "Target Name (only used when updating; blank for new targets)",
        "User Employee Code — maps to User.Employee_Code__c",
        "Criteria Name",
        "Focus Pack Name (blank for non-focus-pack criteria)",
        "Sales Channel",
        "Year",
        "Start Date (YYYY-MM-DD) and End Date (YYYY-MM-DD)",
        "Target Value",
        "Is Active (true / false)",
    ])
    para(doc, "The importer strips BOM, auto-detects , / ; / tab delimiter, "
              "matches headers case-insensitively, resolves users by Employee Code "
              "and criteria / packs by Name. The whole upload is rejected on "
              "validation errors and the failing rows are reported.")

    h(doc, "8.3 Per-row actions", 2)
    bullets(doc, [
        "Recalculate — re-runs the engine just for this target and refreshes the row.",
        "View Calculation — opens a modal with every input the engine used, per-day math, and the final formula. Use this when a user disputes their achievement.",
    ])

    h(doc, "8.4 Export", 2)
    para(doc, "Click Export to write the currently filtered rows to CSV (UTF-8 "
              "with BOM so Excel opens it cleanly). Includes target value, "
              "achievement, %, pending, working days, active flag, last updated.")

    h(doc, "9. Recalculating Achievement", 1)
    para(doc, "Achievement is refreshed by the daily batch. You can also:")
    bullets(doc, [
        "Per target — click Recalculate on a row.",
        "All active targets — click Run Daily Recalc at the top of the Target Manager. Internally this enqueues the SecondaryAchievementBatch for rows whose date range includes today.",
    ])

    h(doc, "10. Configuring Incentive Slabs", 1)
    para(doc, "Open Secondary PBIS Slabs and click New Slab. Slab fields:")
    table(doc,
        ["Field", "Notes"],
        [
            ["Target Criteria",       "The criterion this slab pays out for. Required."],
            ["Sales Channel",         "Slabs are channel-scoped (different channels often pay differently)."],
            ["Focused Pack",          "Optional. When set, beats a generic channel-default slab for the same Focus Pack."],
            ["Compare On",            "Percent (compare Achievement %) or Value (compare Achievement Value)."],
            ["Achievement From / To", "Inclusive bracket. Leave To blank for 'unbounded upwards'."],
            ["Incentive Amount",      "Flat payout when this slab matches."],
            ["Active",                "Inactive slabs are ignored."],
        ],
    )
    para(doc, "Bulk Add Slabs lets you enter many slabs at once for one criterion "
              "+ channel (one row per bracket, with a Clone Row helper).")

    h(doc, "Matching rules", 2)
    para(doc, "For each (User, Criterion, Focus Pack) target row the engine:")
    numbered(doc, [
        "Filters slabs to the user's channel and the target's criterion. Pack-specific slabs are preferred when available.",
        "Picks the value to compare based on Compare On.",
        "Selects all qualifying brackets (From <= value AND (To is null OR value <= To)).",
        "Tie-breaks on highest Achievement_From__c, then highest Incentive_Amount__c.",
        "Writes one Incentive_Credit__c row carrying the criterion, focus pack, matched slab, and snapshot achievement.",
    ])

    h(doc, "11. Running PBIS Monthly", 1)
    para(doc, "Open Secondary PBIS Console.")
    numbered(doc, [
        "Pick Year and Month (defaults to the current month).",
        "Click Run PBIS for this Month. The button is asynchronous (queueable) — it returns immediately with a toast and the screen auto-refreshes after ~5 seconds once the job commits.",
    ])
    para(doc, "Each run is idempotent — prior Incentive_Credit__c rows for that "
              "month are deleted before the fresh rows are inserted, so you can "
              "re-run safely after fixing a slab or a target.")

    h(doc, "What you see after a run", 2)
    table(doc,
        ["Section", "Description"],
        [
            ["Grand Total",        "Sum of all PBIS payouts written for the month."],
            ["Per-user totals",    "One row per user x channel with monthly total and line count. Click View breakdown."],
            ["Breakdown",          "Per-criterion lines for the picked user — Target, Criterion, Focus Pack, Compare On, Achievement, Slab range, Amount."],
            ["Duplicates skipped", "If two active overlapping targets share the same (user, criterion, focus pack) key, the engine deduplicates. Skipped rows and their winners are listed so you can deactivate the duplicates and re-run."],
        ],
    )
    para(doc, "Export downloads every PBIS row for the month as CSV (Executive, "
              "Channel, Year, Month, Target, Criterion, Operator, Focus Pack, "
              "Compare On, Achievement Value, Achievement %, Slab From / To, "
              "Incentive Amount, Computed At).")

    h(doc, "12. Scheduling the Batches", 1)
    para(doc, "Use Setup → Apex Classes → Schedule Apex (or anonymous Apex):")
    table(doc,
        ["Job", "Class", "Suggested cron"],
        [
            ["Daily achievement refresh", "SecondaryAchievementScheduler", "every day 01:30"],
            ["Monthly PBIS",              "SecondaryPBISScheduler",        "02:00 on the 1st of every month — processes the prior month"],
        ],
    )

    h(doc, "13. Period Revisions & Lifecycle Rules", 1)
    para(doc, "When a target needs to change mid-period (BRD §3.1.1 / §6), never edit an active row. The correct flow:")
    numbered(doc, [
        "Open the existing Secondary_Target__c row.",
        "Deactivate it (Is_Active__c = false) and set End_Date__c to the close-out date.",
        "Create a new row for the new effective period with the revised target.",
    ])
    para(doc, "The engine only reads Is_Active__c + the date range, so the closed "
              "row is ignored from then on while remaining as an audit record. The "
              "achievement computed against it is preserved.")

    h(doc, "14. Operational Troubleshooting", 1)
    table(doc,
        ["Symptom", "What to do"],
        [
            ["Target shows Achievement = 0 after running the batch",
             "Open View Calculation on the row — the modal lists raw inputs the engine saw. Most often a Date Field or User Field on the criterion is misconfigured, the user has no Daily_Log__c rows for the period, or the filter excludes matching orders."],
            ["Achievement % suddenly jumps",
             "Targets may have been edited mid-period. Use per-row Recalculate to refresh just that row."],
            ["Two targets for the same user & criterion",
             "PBIS engine keeps the better one and lists the other in 'Duplicates skipped' on the Console. Deactivate the loser."],
            ["Need to re-run PBIS after a slab fix",
             "Click Run PBIS for this Month again — the run is idempotent."],
            ["CSV upload rejected with header / encoding errors",
             "The importer auto-detects ,  ;  and tab delimiters and strips BOM. If it still fails, re-save the file as CSV (UTF-8) from Excel."],
            ["User can't see the KPI Dashboard tab",
             "Tab access is DefaultOn for all 49 profiles. If a custom profile is missing it, add a tabVisibilities entry for Secondary_KPI_Dashboard."],
            ["Dashboard shows duplicate rows for users sharing a role",
             "Fixed — Direct Subordinates are grouped by role with a (+N more) label when several users share one role."],
        ],
    )

    # ===== Appendices =====
    h(doc, "Appendix A — Operator Reference", 1)
    table(doc,
        ["Operator", "BRD Parameter", "Required fields on the criterion"],
        [
            ["SUM",                "Cumulative Secondary Revenue", "Object, Field (numeric), Date Field, User Field, optional Filters"],
            ["COUNT",              "(any plain count)",            "Object, Date Field, User Field, Filters"],
            ["COUNT_DISTINCT",     "UBO / ECO",                    "+ Distinct Field (e.g. Account__c or Order__r.Account__c)"],
            ["DAILY_UNIQUE_AVG",   "TC, PC",                       "+ Distinct Field; tick Use Attendance Divisor; for TC also fill Secondary Source Objects"],
            ["DAILY_RATIO_AVG",    "Average TLSD",                 "Numerator Field (line items), Denominator Field (orders); attendance divisor on"],
            ["FOCUS_PACK_REVENUE", "Focus Pack Revenue",           "Focus Pack is set on each target row"],
            ["FOCUS_PACK_ECO",     "Focus Pack ECO",               "Same; ECO is per-order qualifying — SUM(pack-SKU qty) >= Min Qty in one order"],
        ],
    )
    para(doc, "Working-day divisor — controlled by Use_Attendance_Divisor__c. "
              "Working days = dates with Daily_Log__c.Day_Started_Date__c for that "
              "user. Approved full-day leaves with no Start Day are excluded; "
              "telephonic-only days are excluded from the divisor (though their "
              "orders still contribute to the numerator).")

    h(doc, "Appendix B — Field Reference", 1)
    h(doc, "Secondary_Target__c", 2)
    table(doc,
        ["API Name", "Type", "Notes"],
        [
            ["Name",                   "Auto",          "STGT-0000"],
            ["User__c",                "Lookup → User", "The DSM / SSA"],
            ["User_Name__c",           "Text",          "Snapshot for display"],
            ["Sales_Channel__c",       "Picklist",      "Channel value-set"],
            ["Year__c",                "Number",        ""],
            ["Start_Date__c / End_Date__c", "Date",     "Effective period"],
            ["Target_Criteria__c",     "Lookup",        "The criterion"],
            ["Target_Value__c",        "Number",        "The number to hit"],
            ["Achievement_Value__c",   "Number",        "Engine-written"],
            ["Achievement_Percent__c", "Number",        "Engine-written"],
            ["Pending_Target__c",      "Number",        "Target − Achievement"],
            ["Working_Days__c",        "Number",        "Divisor used (audit)"],
            ["Daily_Achievement__c",   "Number",        "Today's increment"],
            ["Focused_Pack__c",        "Lookup",        "Required for FOCUS_PACK_* criteria"],
            ["Is_Active__c",           "Checkbox",      "The lifecycle gate"],
            ["Last_Updated__c",        "DateTime",      "Last batch / recalc"],
        ],
    )
    h(doc, "Incentive_Slab__c (Secondary additions)", 2)
    table(doc,
        ["API Name", "Notes"],
        [
            ["Target_Criteria__c",      "Identifies a Secondary slab (must be set)"],
            ["Sales_Channel__c",        "Channel filter"],
            ["Focused_Pack__c",         "Optional pack override"],
            ["Compare_On__c",           "Percent or Value"],
            ["Achievement_From__c / Achievement_To__c", "Inclusive bracket"],
            ["Incentive_Amount__c",     "Flat payout"],
            ["Active__c",               "Inactive ⇒ ignored"],
        ],
    )
    h(doc, "Incentive_Credit__c (Secondary additions)", 2)
    table(doc,
        ["API Name", "Notes"],
        [
            ["Source__c",            "Always 'Secondary PBIS' for this module"],
            ["Executive__c",         "User who earned it"],
            ["Year__c / Month__c",   "Natural key for monthly idempotency"],
            ["Sales_Channel__c",     ""],
            ["Target_Criteria__c",   "The criterion that paid"],
            ["Focused_Pack__c",      "Set for pack criteria"],
            ["Matched_Slab__c",      "The slab that decided the amount"],
            ["Secondary_Target__c",  "The source achievement row (audit)"],
            ["Achievement_Value__c / Achievement_Percent__c", "Snapshot at compute time"],
            ["Compare_On__c",        "Snapshot — what the slab compared"],
            ["Credit_Amount__c",     "The paid amount"],
            ["Computed_At__c",       "When the row was written"],
        ],
    )

    out = OUT_DIR / "User_Manual_Admin.docx"
    doc.save(out)
    return out


# ============================================================================
# SALES USER (DSM / SSA) MANUAL
# ============================================================================

def build_sales_user():
    doc = Document()
    style(doc)
    add_cover(doc, "Secondary Targets & PBIS", "Sales User (DSM / SSA)")

    h(doc, "1. Welcome", 1)
    para(doc,
        "This guide is for DSM and SSA field staff. You will see your secondary "
        "targets, how much you have achieved against them, and the PBIS "
        "incentive you have earned for the month — all in one place, on your "
        "phone or desktop."
    )
    para(doc, "You have read-only access to the dashboard. You do not edit "
              "targets or slabs — those are maintained by the MIS / admin team.")

    h(doc, "2. Opening the Dashboard", 1)
    bullets(doc, [
        "Mobile (Salesforce app) — open the bottom-nav More menu and tap Secondary KPI Dashboard.",
        "Desktop — App Launcher → Secondary KPI Dashboard.",
    ])

    h(doc, "3. What You See", 1)
    para(doc, "Because you are a DSM / SSA, the dashboard is locked to your own "
              "data — there is no user picker.")

    h(doc, "3.1 Header", 2)
    bullets(doc, [
        "Your name, your role, your Employee Code, and your Sales Channel.",
        "The period being shown (defaults to the current calendar month).",
    ])

    h(doc, "3.2 Hero tiles", 2)
    table(doc,
        ["Tile", "What it shows"],
        [
            ["Active Targets", "How many secondary targets you have running this period."],
            ["Total Target",   "Sum of your target values across criteria."],
            ["Achievement",    "Sum of what you have achieved, plus your overall Ach %  with a coloured bar."],
            ["PBIS Incentive", "What you have earned for the period, plus the count of credit lines."],
        ],
    )
    para(doc, "Achievement % colour scheme:")
    bullets(doc, [
        ">= 100% — green: you are on or above target.",
        "80% to 99% — amber: close but pushing.",
        "< 80% — red: behind schedule.",
    ])

    h(doc, "3.3 Active Secondary Targets", 2)
    para(doc, "A list of every active target you have. On mobile each target is "
              "a card; on desktop it is a table. Each row shows:")
    table(doc,
        ["Field", "What it means"],
        [
            ["Criterion / Focus Pack / Channel", "What this target is measuring."],
            ["Target",        "The number you have to hit."],
            ["Achievement",   "What you have achieved so far."],
            ["Ach %",         "Percent of target, with the same green / amber / red bucket."],
            ["Pending",       "Target − Achievement (how much further to go)."],
            ["Working Days",  "Working days used in the divisor (for averaging criteria)."],
            ["Incentive",     "The PBIS amount you have earned on this specific criterion this month."],
        ],
    )

    h(doc, "4. Changing the Period", 1)
    para(doc, "Use Year and Month at the top of the dashboard. The dashboard "
              "refreshes immediately when you change either filter.")

    h(doc, "5. Things to Know", 1)
    bullets(doc, [
        "Achievement updates daily. If the latest order isn't reflected yet, wait for the next day's batch — or ask your manager to trigger a per-target recalculation.",
        "Working days affect averages. Days you did not Start Day in the app (no Daily Log) are excluded from the divisor for TC / PC / TLSD-style criteria. Approved leave days are also excluded.",
        "Incentives appear after the monthly PBIS run. The PBIS run typically happens shortly after month-end. Until then, the Incentive tile shows ₹0.",
        "You cannot edit targets or slabs. If a target looks wrong, escalate to your TSE / ASM or the MIS team.",
    ])

    h(doc, "6. Frequently Asked Questions", 1)
    para(doc, "Q. Will achievement update in real time?", bold=True)
    para(doc, "No. It is computed by a daily batch (and an on-demand recalc the "
              "admin can trigger). Real-time recompute on every order write "
              "would be too expensive.")
    para(doc, "Q. Why is my incentive ₹0 even though my achievement is good?", bold=True)
    para(doc, "Incentives are written by the monthly PBIS run, which happens "
              "shortly after month-end. Until then, the Incentive tile shows ₹0.")
    para(doc, "Q. I achieved 150% — am I really 1.5× the target?", bold=True)
    para(doc, "Yes. There is no cap on achievement %. Your slab payout is whatever the highest qualifying bracket pays.")
    para(doc, "Q. I do not see the dashboard tab on my phone.", bold=True)
    para(doc, "Ask your admin to confirm the Secondary KPI Dashboard tab is "
              "added to the mobile navigation for your app.")

    out = OUT_DIR / "User_Manual_Sales_User.docx"
    doc.save(out)
    return out


# ============================================================================
# HIERARCHY USER (TSE / ASM / RSM / Heads) MANUAL
# ============================================================================

def build_hierarchy_user():
    doc = Document()
    style(doc)
    add_cover(doc, "Secondary Targets & PBIS", "Hierarchy User (TSE / ASM / RSM / Heads)")

    h(doc, "1. Welcome", 1)
    para(doc,
        "This guide is for managers and heads who monitor the secondary "
        "performance of every DSM / SSA in their reporting hierarchy. You will "
        "see team-wide analytics, identify top performers and those who need "
        "attention, and drill into any specific DSM / SSA on demand."
    )
    para(doc, "You have read-only access to the dashboard. You do not edit "
              "targets or slabs — those are maintained by the MIS / admin team.")

    h(doc, "2. Opening the Dashboard", 1)
    bullets(doc, [
        "Desktop — App Launcher → Secondary KPI Dashboard.",
        "Mobile (Salesforce app) — bottom-nav More menu → Secondary KPI Dashboard.",
    ])

    h(doc, "3. The Two Modes", 1)
    para(doc, "The dashboard has two modes, switched via the picker at the top:")
    table(doc,
        ["Mode", "When it is shown"],
        [
            ["Team",     "Default. Aggregates and analytics across all your DSM / SSAs."],
            ["Personal", "When you pick a specific DSM / SSA from the filter (or drill into one from the team table)."],
        ],
    )

    h(doc, "4. The Picker (your scope)", 1)
    para(doc, "The 'View' combobox at the top lists:")
    bullets(doc, [
        "All my DSM / SSAs — the default (team mode).",
        "Each individual DSM / SSA in your role-hierarchy downline, labelled 'Name · EMP001 (DSM)'.",
    ])
    para(doc, "You only see DSM / SSAs that report to you (recursively, through "
              "the Salesforce role hierarchy). Finance, distributor and partner "
              "users in your downline do not appear here and do not contaminate "
              "the aggregates. If you have no DSM / SSAs under you, the dashboard "
              "shows a friendly empty-state message.")

    h(doc, "5. Team Mode", 1)

    h(doc, "5.1 Hero tiles", 2)
    table(doc,
        ["Tile", "What it shows"],
        [
            ["Targeted Users", "DSM / SSAs in your team with at least one active secondary target this period."],
            ["Active Targets", "Total across the team."],
            ["Total Target",   "Sum of target values."],
            ["Achievement",    "Sum + overall Ach % with coloured progress bar."],
            ["PBIS Incentive", "Sum of payouts earned by the team this period."],
        ],
    )
    para(doc, "Colour scheme on Ach %:")
    bullets(doc, [
        ">= 100% — green",
        "80% to 99% — amber",
        "< 80% — red",
    ])

    h(doc, "5.2 My DSM / SSAs", 2)
    para(doc, "A table of every team member with secondary activity, sorted by "
              "Ach % descending. Columns: Name, Role, Channel, Targets, Target, "
              "Achievement, Ach %, Incentive.")
    bullets(doc, [
        "Drill in — pick the 'View this user' row action (desktop) or tap the card (mobile). The whole dashboard switches to Personal mode for that user.",
        "A 'Back to team' button appears in personal mode to return.",
    ])

    h(doc, "5.3 Top Performers", 2)
    para(doc, "Top 5 DSM / SSAs by Ach %, requiring a non-zero target so empty "
              "users don't dominate. Useful for monthly recognition.")

    h(doc, "5.4 Needs Attention", 2)
    para(doc, "Bottom 5 by Ach %, same filter. Useful for proactive coaching.")

    h(doc, "5.5 Achievement by Sales Channel", 2)
    para(doc, "One bar per channel (CPD / TN / KN / etc.) showing the channel's "
              "overall Ach %, plus the achievement and target totals and how many "
              "users contributed.")

    h(doc, "5.6 Achievement by Criterion", 2)
    para(doc, "One bar per criterion (Revenue, TC, PC, FP ECO, …) — the same "
              "view from the criterion's perspective. Use it to spot which "
              "parameters the team is under-performing on.")

    h(doc, "6. Personal Mode (picking one DSM / SSA)", 1)
    para(doc, "When you select a specific user from the picker — or drill in "
              "from the team table — the dashboard becomes the same view a "
              "DSM / SSA sees of themselves.")
    bullets(doc, [
        "Hero with that user's totals.",
        "Per-target list with Criterion, Focus Pack, Channel, Target, Achievement, Pending, Incentive.",
    ])
    para(doc, "Use the 'Back to team' button (top right) to return to team mode.")

    h(doc, "7. Things to Know", 1)
    bullets(doc, [
        "The picker filters to DSM / SSA profiles only — finance, distributor, partner-community users in your downline do not appear, and they do not contaminate the aggregates.",
        "You can only view DSM / SSAs inside your hierarchy. If a user id outside your hierarchy is somehow passed (URL hack, etc.) the server falls back to team mode — no data leak.",
        "DSM / SSAs cannot see this view. When a DSM / SSA logs in, the picker is hidden and the dashboard is locked to themselves on the server, regardless of any client-side state.",
        "Period filter (Year + Month) applies to the whole dashboard.",
        "Achievement updates daily; incentives appear after the monthly PBIS run.",
    ])

    h(doc, "8. Frequently Asked Questions", 1)
    para(doc, "Q. A DSM is showing 1585% achievement — is that real?", bold=True)
    para(doc, "There is no cap on achievement %. If the target was set low or "
              "the achievement is genuinely above target, the percent will be "
              "high. Verify the target value was set correctly in the Secondary "
              "Target Manager.")
    para(doc, "Q. Some users are missing from 'My DSM / SSAs'.", bold=True)
    para(doc, "The table only shows users with secondary activity in the "
              "selected month (at least one active target or one incentive line). "
              "A DSM / SSA with no active target for this period will not appear "
              "in the team list (but is still visible in the picker for the "
              "current and future months).")
    para(doc, "Q. Why are Distributors or Finance users not listed?", bold=True)
    para(doc, "By design — the dashboard is scoped to DSM / SSA only. Other "
              "profiles are filtered out of both the picker and the analytics.")
    para(doc, "Q. The 'Back to team' button is missing.", bold=True)
    para(doc, "You are either already in team mode, or you are logged in as a "
              "DSM / SSA (who is locked to their own personal view by design).")

    out = OUT_DIR / "User_Manual_Hierarchy_User.docx"
    doc.save(out)
    return out


def main():
    paths = []
    paths.append(build_admin())
    paths.append(build_sales_user())
    paths.append(build_hierarchy_user())
    print("Generated:")
    for p in paths:
        print(f"  {p}")


if __name__ == "__main__":
    main()
